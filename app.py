import os, time, sqlite3, uuid
import streamlit as st
import yaml
import pandas as pd
from typing import Dict, Any, Optional, List


# Initialize session state for page navigation and input handling
if "current_page" not in st.session_state:
    st.session_state.current_page = "main"
if "input_state" not in st.session_state:
    # States: INITIAL, AWAITING_INPUT, INPUT_COLLECTED, PROCESSING, COMPLETED, ERROR
    st.session_state.input_state = "INITIAL"
if "collected_inputs" not in st.session_state:
    st.session_state.collected_inputs = {}
if "required_inputs_config" not in st.session_state:
    st.session_state.required_inputs_config = []
if "current_input_index" not in st.session_state:
    st.session_state.current_input_index = 0
if "human_input_request" not in st.session_state:
    st.session_state.human_input_request = None # Will store {'id':uuid, 'question':'...', 'response':None, 'ready':False}
if "crew_running" not in st.session_state:
    st.session_state.crew_running = False
if "proposal_id" not in st.session_state:
    st.session_state.proposal_id = None
if "chat_messages" not in st.session_state:
    st.session_state.chat_messages = [] # Initialize chat messages


# Conditional page config only if on main page initially
if st.session_state.current_page == "main" and 'page_config_set' not in st.session_state:
    st.set_page_config(page_title="🎯 BidMaster", layout="wide")
    st.session_state.page_config_set = True # Mark config as set


import proposal_team # Import your crew setup
from proposal_tools import parse_store_input_docs, store_proposal_section, get_proposal_sections

# --- Configuration and DB Setup ---
CONFIG_DIR = "config"
DB_PATH = "proposals.db" # Defined in proposal_tools

def load_yaml_config(file_path):
    try:
        with open(file_path, "r") as file:
            return yaml.safe_load(file)
    except FileNotFoundError:
        st.error(f"Error: Configuration file not found at {file_path}")
        return None
    except yaml.YAMLError as e:
        st.error(f"Error parsing YAML file {file_path}: {e}")
        return None

def save_yaml_config(file_path, data):
     with open(file_path, "w") as file:
         yaml.dump(data, file, default_flow_style=False)

# Load required inputs configuration
inputs_config_data = load_yaml_config(os.path.join(CONFIG_DIR, "inputs.yaml"))
if inputs_config_data:
    st.session_state.required_inputs_config = inputs_config_data.get("required_inputs", [])


# --- UI Functions ---
def show_section_modal(section_head, content):
    # ... (keep existing function)
    """Displays the proposal section in a modal."""
    section_placeholder = st.empty()
    with section_placeholder.container():
        with st.expander(f"📄 {section_head}", expanded=False):
            st.markdown(content, unsafe_allow_html=True)

def format_and_download_section_docx(section_name, content, proposal_id, logo_path=None):
    """Formats the section output into a DOCX file and provides a download."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
    from io import BytesIO

    CORPORATE_HEADER = "Acme Corp Proposal"
    CORPORATE_FOOTER = "Confidential"

    document = Document()
    header = document.sections[0].header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = CORPORATE_HEADER
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    section_paragraph = document.add_paragraph()
    section_paragraph.text = f"Section: {section_name.capitalize()}"
    section_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # section_paragraph.style.font.size = Pt(14) # Causing error, commented out

    document.add_paragraph(content)  

    footer = document.sections[0].footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = CORPORATE_FOOTER
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    filename = f"{proposal_id}_{section_name}.docx"
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    st.download_button(
        label=f"⬇️ Download {section_name.capitalize()} (DOCX)",
        data=buffer,
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key=f"download_docx_{proposal_id}_{section_name}" # Unique key
    )


def display_sections(proposal_id):
    if not proposal_id:
        st.info("Select or start a proposal to see generated outputs.")
        return

    sections = get_proposal_sections(proposal_id) # Pass proposal_id

    if not sections:
        st.info("Crew generated proposal outputs will appear here...")
    else:
        for phase_name, phase_data in sections.items():
            phase_output = phase_data.get("phase_output", "No output available.")
            show_section_modal(phase_name.capitalize() + " Output", phase_output)
            # Ensure proposal_id is passed here correctly
            format_and_download_section_docx(phase_name, phase_output, proposal_id)

            if phase_data.get("tasks"):
                with st.expander(f"Detailed Tasks for {phase_name.capitalize()}", expanded=False):
                    for task in phase_data["tasks"]:
                        st.markdown(f"**Task:** {task['task_name'].capitalize()}")
                        st.markdown(f"**Output:** {task['output']}", unsafe_allow_html=True)
                        st.caption(f"Created at: {task['created_at']}")
                        st.divider()

# Function to run crews sequentially
def run_proposal_crews(inputs):
    st.session_state.crew_running = True
    log_placeholder = st.empty() # Placeholder for log expander

    try:
        with log_placeholder.container():
            with st.expander("📜 Crew Execution Log", expanded=True):
                st.info("Starting Analysis Crew...")
                st.session_state["processing_phase"] = "analysis" # Set phase for callback
                analysis_output = proposal_team.proposal_analysis_crew.kickoff(inputs=inputs)
                st.success("✅ Analysis Crew Completed.")
                st.write("Output:", analysis_output) # Display raw output or summary
                # Optionally rename log file here if output_log_file is used
                # if os.path.exists("./logs/analysis.txt"):
                #     os.rename("./logs/analysis.txt", f"./logs/{inputs['proposal_id']}_analysis.log")


                # Check if human input was requested during the crew run
                # (This check might be needed after each kickoff if tool can be called anytime)
                if st.session_state.get('human_input_request') and not st.session_state.human_input_request.get('ready'):
                     st.warning("Crew paused, waiting for human input in chat.")
                     st.session_state.crew_running = False # Pause execution flow
                     st.rerun() # Rerun to update chat

                st.info("Starting Solution Crew...")
                st.session_state["processing_phase"] = "approach"
                solution_output = proposal_team.proposal_solution_crew.kickoff(inputs=inputs)
                st.success("✅ Solution Crew Completed.")
                st.write("Output:", solution_output)
                # if os.path.exists("./logs/approach.txt"):
                #      os.rename("./logs/approach.txt", f"./logs/{inputs['proposal_id']}_approach.log")

                if st.session_state.get('human_input_request') and not st.session_state.human_input_request.get('ready'):
                     st.warning("Crew paused, waiting for human input in chat.")
                     st.session_state.crew_running = False
                     st.rerun()

                st.info("Starting Drafting Crew...")
                st.session_state["processing_phase"] = "draft"
                drafting_output = proposal_team.proposal_drafting_crew.kickoff(inputs=inputs)
                st.success("✅ Drafting Crew Completed.")
                st.write("Output:", drafting_output)
                # if os.path.exists("./logs/draft.txt"):
                #      os.rename("./logs/draft.txt", f"./logs/{inputs['proposal_id']}_draft.log")

                st.session_state.input_state = "COMPLETED"
                st.session_state["processing_phase"] = None # Reset phase
                st.balloons()

    except Exception as e:
        st.error(f"❌ An error occurred during crew execution: {e}")
        st.session_state.input_state = "ERROR"
    finally:
        st.session_state.crew_running = False
        # Option to clear the log expander or keep it
        # log_placeholder.empty()
        st.info("Processing finished. Refreshing sections...")
        time.sleep(2) # Give a moment before rerun
        st.rerun() # Rerun to update displayed sections

# Function to create a new proposal record
def create_proposal_master(proposal_name):
    proposal_id = str(uuid.uuid4())
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("INSERT INTO proposal_master (proposal_id, proposal_name) VALUES (?, ?)", (proposal_id, proposal_name))
    conn.commit()
    conn.close()
    return proposal_id    

# --- Main App Logic ---

if st.session_state.current_page == "main":

    st.markdown("<h1 style='text-align: center;'>🎯 BidMaster</h1>", unsafe_allow_html=True)

    # ✅ Sidebar - Proposal Selection
    st.sidebar.header("Proposal Management")
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT proposal_id, proposal_name FROM proposal_master ORDER BY created_at DESC")
    proposals = cursor.fetchall()
    conn.close()

    # Create the list of options in the desired order
    selectbox_options = ["🆕 Start New Proposal"]
    # Create the mapping dictionary
    proposal_options_map = {"🆕 Start New Proposal": "NEW"}

    # Add existing proposals to the list and map
    for pid, name in proposals:
        display_name = f"{name} ({pid[:8]})" # Format for display
        selectbox_options.append(display_name)
        proposal_options_map[display_name] = pid

    # Use the pre-ordered list for the selectbox options
    selected_option = st.sidebar.selectbox(
        "Select Proposal or Start New:",
        options=selectbox_options, # Use the ordered list here
        index=0 # Ensure "Start New" is selected by default
    )

    # Get the corresponding ID from the map
    selected_proposal_id = proposal_options_map[selected_option]
    
    proposal_name_input = ""
    if selected_proposal_id == "NEW":
        proposal_name_input = st.sidebar.text_input("Enter New Proposal Name:", key="new_proposal_name")
        st.session_state.proposal_id = None # Clear proposal ID if starting new
        if st.session_state.input_state != "INITIAL":
             # Reset state if user switches back to 'new' after starting input/processing
             st.session_state.input_state = "INITIAL"
             st.session_state.collected_inputs = {}
             st.session_state.current_input_index = 0
             st.session_state.chat_messages = [] # Clear chat too
    else:
        if st.session_state.proposal_id != selected_proposal_id:
             # If selection changes, load its state
             st.session_state.proposal_id = selected_proposal_id
             st.session_state.input_state = "INPUT_COLLECTED" # Assume existing proposal has inputs
             st.session_state.collected_inputs = {} # Ideally load saved inputs if stored
             st.session_state.chat_messages = [] # Load chat history from DB below
             st.session_state.current_input_index = len(st.session_state.required_inputs_config) # Mark inputs as done
             st.rerun() # Rerun to load data for the selected proposal


    # File Uploader
    st.sidebar.subheader("📂 Upload Documents")
    uploaded_files = st.sidebar.file_uploader(
        "Upload RFP/related documents (PDF/DOCX)",
        type=["pdf", "docx"],
        accept_multiple_files=True,
        key="file_uploader"
    )

    # Process Files Button
    saved_paths = []
    if uploaded_files:
         if selected_proposal_id == "NEW" and not proposal_name_input.strip():
              st.sidebar.warning("Please enter a name for the new proposal.")
         elif selected_proposal_id == "NEW" and proposal_name_input.strip():
              if st.sidebar.button("Create Proposal & Process Files", key="create_process"):
                  st.session_state.proposal_id = create_proposal_master(proposal_name_input) # Use function from proposal_team
                  st.session_state.input_state = "INITIAL" # Ready to collect inputs
                  st.session_state.chat_messages = [] # Clear chat for new proposal
                  upload_dir = f"uploads/{st.session_state.proposal_id}/"
                  os.makedirs(upload_dir, exist_ok=True)
                  for file in uploaded_files:
                       file_path = os.path.join(upload_dir, file.name)
                       with open(file_path, "wb") as f:
                            f.write(file.getbuffer())
                       saved_paths.append(file_path)
                  parse_store_input_docs(st.session_state.proposal_id, saved_paths) # Process files
                  st.sidebar.success(f"Proposal '{proposal_name_input}' created (ID: ...{st.session_state.proposal_id[-6:]}). Files processed.")
                  st.session_state.input_state = "AWAITING_INPUT" # Start input collection
                  st.session_state.current_input_index = 0
                  st.session_state.collected_inputs = {'proposal_id': st.session_state.proposal_id, 'document_list': saved_paths} # Store initial inputs
                  #st.rerun()
         elif selected_proposal_id != "NEW":
              if st.sidebar.button("Process Files for Selected Proposal", key="process_existing"):
                  upload_dir = f"uploads/{st.session_state.proposal_id}/"
                  os.makedirs(upload_dir, exist_ok=True)
                  for file in uploaded_files:
                       file_path = os.path.join(upload_dir, file.name)
                       with open(file_path, "wb") as f:
                            f.write(file.getbuffer())
                       saved_paths.append(file_path)
                  parse_store_input_docs(st.session_state.proposal_id, saved_paths) # Process files
                  st.sidebar.success(f"Files processed for proposal ID: ...{st.session_state.proposal_id[-6:]}.")
                  # Update document list if kickoff happens later
                  if 'document_list' in st.session_state.collected_inputs:
                       st.session_state.collected_inputs['document_list'].extend(saved_paths)
                  else:
                       st.session_state.collected_inputs['document_list'] = saved_paths
    print(f"Current Session State: {st.session_state.input_state}")


    # Start Generation Button / Input Collection Trigger
    st.sidebar.subheader("🚀 Generate Proposal")
    can_start_generation = (st.session_state.proposal_id is not None and
                            st.session_state.input_state == "INPUT_COLLECTED" and
                            not st.session_state.crew_running)

    if st.sidebar.button("Start Generation", disabled=not can_start_generation, key="start_gen"):
         st.session_state.input_state = "PROCESSING"
         # Prepare final inputs for kickoff
         final_inputs = st.session_state.collected_inputs.copy()
         # Ensure proposal_id and document_list are present
         if 'proposal_id' not in final_inputs:
              final_inputs['proposal_id'] = st.session_state.proposal_id
         if 'document_list' not in final_inputs:
              # Attempt to find saved paths if not processed in this session
              upload_dir = f"uploads/{st.session_state.proposal_id}/"
              if os.path.exists(upload_dir):
                   final_inputs['document_list'] = [os.path.join(upload_dir, f) for f in os.listdir(upload_dir)]
              else:
                   final_inputs['document_list'] = []

         st.info("🚀 Kicking off proposal generation crews...")
         # Use threading or similar for non-blocking execution if needed
         # For simplicity, running directly here
         run_proposal_crews(final_inputs)
         st.rerun()

    # --- Configuration Button ---
    if st.sidebar.button("⚙️ Open Configuration", key="open_config"):
        st.session_state.current_page = "config"
        st.rerun()

    # --- Main Page UI ---
    st.subheader("Proposal Hub", divider=True)

    # Display Generated Sections
    st.subheader("📄 Generated outputs")
    display_sections(st.session_state.proposal_id) # Pass current proposal ID

    # Chat UI
    st.subheader("💬 Message Board")

    # Load chat history only when proposal ID changes or initially
    if 'last_loaded_proposal_id_chat' not in st.session_state or st.session_state.last_loaded_proposal_id_chat != st.session_state.proposal_id:
        if st.session_state.proposal_id:
            # Load historical messages from DB for the selected proposal
            # Assume get_chat_history function exists and works
            # st.session_state.chat_messages = get_chat_history(st.session_state.proposal_id)
            st.session_state.chat_messages = [] # Start fresh or load from DB
            st.session_state.last_loaded_proposal_id_chat = st.session_state.proposal_id
        else:
            st.session_state.chat_messages = []
            st.session_state.last_loaded_proposal_id_chat = None


    # Display chat messages from session state
    for message in st.session_state.chat_messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    # Display prompt for next required input if awaiting input
    if st.session_state.input_state == "AWAITING_INPUT" and st.session_state.current_input_index < len(st.session_state.required_inputs_config):
        current_input_cfg = st.session_state.required_inputs_config[st.session_state.current_input_index]
        input_prompt = current_input_cfg['prompt']
        # Display the prompt only once per state change
        if 'last_displayed_prompt' not in st.session_state or st.session_state.last_displayed_prompt != input_prompt:
             with st.chat_message("assistant"):
                  st.markdown(f"Please provide the following information:\n\n**{input_prompt}**")
             st.session_state.last_displayed_prompt = input_prompt
             st.session_state.chat_messages.append({"role": "assistant", "content": f"Please provide the following information:\n\n**{input_prompt}**"})


    # Display prompt for human feedback tool if requested
    human_request = st.session_state.get('human_input_request')
    if human_request and not human_request.get('ready'):
        question = human_request['question']
        if 'last_displayed_human_question_id' not in st.session_state or st.session_state.last_displayed_human_question_id != human_request['id']:
             with st.chat_message("assistant"):
                  st.warning(f"⚠️ Human input needed by the crew:")
                  st.markdown(question)
             st.session_state.last_displayed_human_question_id = human_request['id']
             st.session_state.chat_messages.append({"role": "assistant", "content": f"⚠️ **Human input needed:** {question}"})


    # Chat Input Handling
    user_input = st.chat_input("Send a message or provide requested input...")
    if user_input:
        print(f"Current Session State: {st.session_state.input_state}")
 
        # Add user message to chat display
        st.session_state.chat_messages.append({"role": "user", "content": user_input})
    
        # Optionally store user message to DB
        # if st.session_state.proposal_id:
        #     store_chat_message(st.session_state.proposal_id, "user", user_input)

        # --- Input Collection Logic ---
        if st.session_state.input_state == "AWAITING_INPUT":
            if st.session_state.current_input_index < len(st.session_state.required_inputs_config):
                current_input_cfg = st.session_state.required_inputs_config[st.session_state.current_input_index]
                input_id = current_input_cfg['id']
                st.session_state.collected_inputs[input_id] = user_input # Store collected input
                st.session_state.current_input_index += 1 # Move to next input

                # Check if all inputs are collected
                if st.session_state.current_input_index >= len(st.session_state.required_inputs_config):
                    st.session_state.input_state = "INPUT_COLLECTED"
                    st.session_state.chat_messages.append({"role": "assistant", "content": "✅ All required inputs collected. You can now click 'Start Generation'."})
                    del st.session_state.last_displayed_prompt # Clear prompt tracking
                else:
                    # Prepare for the next input prompt in the next rerun cycle
                    del st.session_state.last_displayed_prompt
            st.rerun() # Rerun to display next prompt or confirmation

        # --- Human Feedback Tool Response Logic ---
        elif st.session_state.get('human_input_request') and not st.session_state['human_input_request'].get('ready'):
            # This input is the response to the agent's question
            st.session_state['human_input_request']['response'] = user_input
            st.session_state['human_input_request']['ready'] = True # Signal the tool to proceed
            del st.session_state.last_displayed_human_question_id # Clear prompt tracking
            st.session_state.chat_messages.append({"role": "assistant", "content": f"✅ Got it. Providing response to the crew."})

            # If the crew was paused, maybe trigger a rerun to potentially resume it
            # This depends on how run_proposal_crews handles the pause
            if not st.session_state.crew_running:
                 st.warning("Attempting to resume crew... (Manual restart might be needed if it doesn't proceed)")
                 # Potentially re-call run_proposal_crews if state allows, or just let the tool unblock
            st.rerun()

        # --- General Chat / Feedback Logic (Optional) ---
        else:
            print("Inside Else")
            # Handle general chat or feedback if needed
            # Example: Check for keywords like "feedback"
            if "feedback" in user_input.lower() and st.session_state.proposal_id:
                 # Add logic to process feedback if required, potentially kicking off specific tasks
                 st.session_state.chat_messages.append({"role": "assistant", "content": "🔄 Feedback received. Processing..."})
                 # ... (Add feedback processing logic here) ...
            st.rerun() # Rerun to show the user's message immediately
        print(f"Outside Else")  # Debug
        print(f"Current input index: {st.session_state.current_input_index}")  # Debug
        print(f"Required inputs config: {st.session_state.required_inputs_config}")  # Debug
        print(f"Collected inputs: {st.session_state.collected_inputs}")  # Debug       

# --- Configuration Page Logic ---
elif st.session_state.current_page == "config":
    # Set page config here if navigating directly to config (might cause issues if already set)
    if 'page_config_set' not in st.session_state:
         st.set_page_config(page_title="⚙️ Configuration", layout="wide")
         st.session_state.page_config_set = True

    if st.button("❌ Close Configuration", key="close_config"):
        st.session_state.current_page = "main"
        del st.session_state.page_config_set # Allow main page to set config again
        st.rerun()

    st.title("⚙️ Configuration Panel")

    # Load fresh configs when entering the page
    config_files = {
        "Agents": os.path.join(CONFIG_DIR, "agents.yaml"),
        "Tasks": os.path.join(CONFIG_DIR, "tasks.yaml"),
        "LLM": os.path.join(CONFIG_DIR, "llm.yaml"),
        "Inputs": os.path.join(CONFIG_DIR, "inputs.yaml"), # Add new inputs file
        # "External Tools": os.path.join(CONFIG_DIR, "tools.yaml"), # Uncomment if you have this
    }
    configs = {}
    for name, path in config_files.items():
         loaded_config = load_yaml_config(path)
         if loaded_config is not None:
              configs[name] = loaded_config
         else:
              # Handle case where a config file failed to load
              st.error(f"Failed to load configuration for {name}")
              configs[name] = {} # Provide an empty dict to avoid errors

    tabs = st.tabs(list(configs.keys()))

    for i, name in enumerate(configs.keys()):
        path = config_files[name]
        current_config_data = configs[name]

        with tabs[i]:
            st.subheader(f"{name} Configuration (`{os.path.basename(path)}`)")

            # Display YAML content in an editable text area
            try:
                yaml_str = yaml.dump(current_config_data, default_flow_style=False, sort_keys=False)
            except Exception as e:
                st.error(f"Error converting {name} config to YAML string: {e}")
                yaml_str = "# Error displaying YAML"

            edited_yaml_str = st.text_area(f"Edit {name} YAML:", value=yaml_str, height=400, key=f"yaml_edit_{name}")

            # Save Button
            if st.button(f"Save {name} Config", key=f"save_{name}"):
                try:
                    new_config_data = yaml.safe_load(edited_yaml_str)
                    save_yaml_config(path, new_config_data)
                    configs[name] = new_config_data # Update in-memory config
                    # Update relevant session state if needed (e.g., inputs config)
                    if name == "Inputs":
                        st.session_state.required_inputs_config = new_config_data.get("required_inputs", [])
                    st.success(f"{name} configuration saved successfully!")
                    time.sleep(1)
                    st.rerun() # Rerun to reflect changes if needed elsewhere
                except yaml.YAMLError as e:
                    st.error(f"Invalid YAML format in {name}: {e}")
                except Exception as e:
                    st.error(f"Error saving {name} config: {e}")